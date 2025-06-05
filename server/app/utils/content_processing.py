"""Content processing utilities for AI question generation."""

import os
import re
import json
import mimetypes
import tempfile
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from urllib.parse import urlparse
import logging

# Import content processing libraries with fallbacks
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import whisper
    HAS_WHISPER = True
except ImportError:
    HAS_WHISPER = False

try:
    import requests
    from bs4 import BeautifulSoup
    HAS_WEB = True
except ImportError:
    HAS_WEB = False

try:
    import nltk
    from textstat import flesch_reading_ease, flesch_kincaid_grade
    HAS_TEXTSTAT = True
except ImportError:
    HAS_TEXTSTAT = False

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.decomposition import LatentDirichletAllocation
    HAS_SKLEARN = True
except ImportError:
    HAS_SKLEARN = False

try:
    import spacy
    HAS_SPACY = True
except ImportError:
    HAS_SPACY = False


logger = logging.getLogger(__name__)


class ContentExtractor:
    """Extract content from various file formats."""
    
    def __init__(self):
        self.supported_formats = {
            'pdf': self.extract_pdf,
            'docx': self.extract_docx,
            'doc': self.extract_docx,  # Fallback
            'txt': self.extract_text,
            'md': self.extract_text,
            'rtf': self.extract_text,
            'mp3': self.extract_audio,
            'wav': self.extract_audio,
            'm4a': self.extract_audio,
            'mp4': self.extract_video,
            'avi': self.extract_video,
            'mov': self.extract_video,
            'url': self.extract_url
        }
    
    def extract_content(self, source_path: str, content_type: str = None) -> Dict[str, Any]:
        """Extract content from various sources."""
        try:
            if content_type == 'url' or source_path.startswith(('http://', 'https://')):
                return self.extract_url(source_path)
            
            if not os.path.exists(source_path):
                return self._error_result(f"File not found: {source_path}")
            
            # Determine file type
            file_ext = Path(source_path).suffix.lower().lstrip('.')
            if not file_ext:
                # Try to detect by MIME type
                mime_type, _ = mimetypes.guess_type(source_path)
                if mime_type:
                    file_ext = mime_type.split('/')[-1]
            
            if file_ext not in self.supported_formats:
                return self._error_result(f"Unsupported file format: {file_ext}")
            
            # Extract content using appropriate method
            extractor = self.supported_formats[file_ext]
            result = extractor(source_path)
            
            # Add file metadata
            if result['success']:
                result['metadata'].update({
                    'file_path': source_path,
                    'file_size': os.path.getsize(source_path),
                    'file_type': file_ext,
                    'extraction_method': extractor.__name__
                })
            
            return result
            
        except Exception as e:
            logger.error(f"Content extraction failed: {str(e)}")
            return self._error_result(str(e))
    
    def extract_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF files."""
        if not HAS_PDF:
            return self._error_result("PDF processing not available. Install PyPDF2.")
        
        try:
            text = ""
            metadata = {'pages': 0, 'title': '', 'author': ''}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata['pages'] = len(pdf_reader.pages)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata['title'] = pdf_reader.metadata.get('/Title', '')
                    metadata['author'] = pdf_reader.metadata.get('/Author', '')
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num + 1}: {str(e)}")
            
            if not text.strip():
                return self._error_result("No text content found in PDF")
            
            return self._success_result(text, metadata)
            
        except Exception as e:
            return self._error_result(f"PDF extraction failed: {str(e)}")
    
    def extract_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX files."""
        if not HAS_DOCX:
            return self._error_result("DOCX processing not available. Install python-docx.")
        
        try:
            doc = docx.Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(cell for cell in row_data):
                        table_data.append(" | ".join(row_data))
                if table_data:
                    tables_text.append("\n".join(table_data))
            
            # Combine all text
            text = "\n\n".join(paragraphs)
            if tables_text:
                text += "\n\n--- Tables ---\n" + "\n\n".join(tables_text)
            
            if not text.strip():
                return self._error_result("No text content found in document")
            
            metadata = {
                'paragraphs': len(paragraphs),
                'tables': len(tables_text),
                'has_images': len(doc.inline_shapes) > 0
            }
            
            # Extract core properties if available
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                metadata.update({
                    'title': props.title or '',
                    'author': props.author or '',
                    'subject': props.subject or '',
                    'created': props.created.isoformat() if props.created else '',
                    'modified': props.modified.isoformat() if props.modified else ''
                })
            
            return self._success_result(text, metadata)
            
        except Exception as e:
            return self._error_result(f"DOCX extraction failed: {str(e)}")
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from plain text files."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                return self._error_result("Could not decode text file with any supported encoding")
            
            if not text.strip():
                return self._error_result("File is empty")
            
            metadata = {
                'encoding': encoding,
                'lines': len(text.splitlines()),
                'file_format': 'plain_text'
            }
            
            return self._success_result(text, metadata)
            
        except Exception as e:
            return self._error_result(f"Text extraction failed: {str(e)}")
    
    def extract_audio(self, file_path: str) -> Dict[str, Any]:
        """Extract text from audio files using Whisper."""
        if not HAS_WHISPER:
            return self._error_result("Audio processing not available. Install openai-whisper.")
        
        try:
            # Load Whisper model
            model = whisper.load_model("base")
            
            # Transcribe audio
            result = model.transcribe(file_path)
            
            text = result["text"]
            if not text.strip():
                return self._error_result("No speech detected in audio file")
            
            metadata = {
                'duration': result.get('duration', 0),
                'language': result.get('language', 'unknown'),
                'segments': len(result.get('segments', [])),
                'transcription_model': 'whisper-base'
            }
            
            # Add segment information if available
            if 'segments' in result:
                segments_info = []
                for segment in result['segments']:
                    segments_info.append({
                        'start': segment.get('start', 0),
                        'end': segment.get('end', 0),
                        'text': segment.get('text', '').strip()
                    })
                metadata['detailed_segments'] = segments_info
            
            return self._success_result(text, metadata)
            
        except Exception as e:
            return self._error_result(f"Audio extraction failed: {str(e)}")
    
    def extract_video(self, file_path: str) -> Dict[str, Any]:
        """Extract text from video files (audio transcription)."""
        # For video files, we extract the audio and transcribe it
        # In a production environment, you might use ffmpeg to extract audio first
        try:
            # For now, treat video same as audio
            # In production, you would:
            # 1. Extract audio track using ffmpeg
            # 2. Save to temporary file
            # 3. Transcribe the audio
            # 4. Clean up temporary file
            
            return self.extract_audio(file_path)
            
        except Exception as e:
            return self._error_result(f"Video extraction failed: {str(e)}")
    
    def extract_url(self, url: str) -> Dict[str, Any]:
        """Extract text from web pages."""
        if not HAS_WEB:
            return self._error_result("Web content processing not available. Install requests and beautifulsoup4.")
        
        try:
            # Validate URL
            parsed_url = urlparse(url)
            if not parsed_url.scheme or not parsed_url.netloc:
                return self._error_result("Invalid URL format")
            
            # Fetch content
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=30, allow_redirects=True)
            response.raise_for_status()
            
            # Parse HTML
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove unwanted elements
            for element in soup(['script', 'style', 'nav', 'header', 'footer', 'aside']):
                element.decompose()
            
            # Extract title
            title = soup.find('title')
            title_text = title.get_text().strip() if title else ''
            
            # Extract main content
            # Try to find main content areas
            main_content = soup.find('main') or soup.find('article') or soup.find('div', class_=re.compile(r'content|main|article', re.I))
            
            if main_content:
                text_content = main_content.get_text()
            else:
                # Fallback to body content
                text_content = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text_content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            if not text.strip():
                return self._error_result("No text content found on webpage")
            
            metadata = {
                'url': url,
                'title': title_text,
                'content_type': response.headers.get('content-type', ''),
                'content_length': len(response.content),
                'status_code': response.status_code,
                'final_url': response.url,  # In case of redirects
                'extraction_method': 'beautifulsoup'
            }
            
            # Extract meta description
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            if meta_desc:
                metadata['description'] = meta_desc.get('content', '')
            
            # Extract meta keywords
            meta_keywords = soup.find('meta', attrs={'name': 'keywords'})
            if meta_keywords:
                metadata['keywords'] = meta_keywords.get('content', '')
            
            return self._success_result(text, metadata)
            
        except requests.exceptions.RequestException as e:
            return self._error_result(f"Failed to fetch URL: {str(e)}")
        except Exception as e:
            return self._error_result(f"Web content extraction failed: {str(e)}")
    
    def _success_result(self, text: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Create success result."""
        return {
            'success': True,
            'text': text,
            'metadata': metadata,
            'word_count': len(text.split()),
            'character_count': len(text),
            'error': None
        }
    
    def _error_result(self, error_message: str) -> Dict[str, Any]:
        """Create error result."""
        return {
            'success': False,
            'text': '',
            'metadata': {},
            'word_count': 0,
            'character_count': 0,
            'error': error_message
        }


class TextAnalyzer:
    """Analyze text content for various metrics."""
    
    def __init__(self):
        self.nlp = None
        if HAS_SPACY:
            try:
                import spacy
                self.nlp = spacy.load("en_core_web_sm")
            except OSError:
                logger.warning("SpaCy English model not found. Some features will be limited.")
    
    def analyze_text(self, text: str) -> Dict[str, Any]:
        """Perform comprehensive text analysis."""
        try:
            analysis = {
                'readability': self._analyze_readability(text),
                'complexity': self._analyze_complexity(text),
                'keywords': self._extract_keywords(text),
                'topics': self._extract_topics(text),
                'language_features': self._analyze_language_features(text),
                'structure': self._analyze_structure(text)
            }
            
            # Calculate overall difficulty score (1-10)
            difficulty_score = self._calculate_difficulty_score(analysis)
            analysis['difficulty_score'] = difficulty_score
            
            return analysis
            
        except Exception as e:
            logger.error(f"Text analysis failed: {str(e)}")
            return {
                'readability': {},
                'complexity': {},
                'keywords': [],
                'topics': [],
                'language_features': {},
                'structure': {},
                'difficulty_score': 5.0,
                'error': str(e)
            }
    
    def _analyze_readability(self, text: str) -> Dict[str, Any]:
        """Analyze text readability."""
        readability = {}
        
        if HAS_TEXTSTAT:
            try:
                readability = {
                    'flesch_reading_ease': flesch_reading_ease(text),
                    'flesch_kincaid_grade': flesch_kincaid_grade(text),
                    'reading_level': self._get_reading_level(flesch_reading_ease(text))
                }
            except Exception as e:
                logger.warning(f"Readability analysis failed: {str(e)}")
        
        # Basic readability metrics
        words = text.split()
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        if words and sentences:
            avg_words_per_sentence = len(words) / len(sentences)
            avg_syllables_per_word = sum(self._count_syllables(word) for word in words) / len(words)
            
            readability.update({
                'word_count': len(words),
                'sentence_count': len(sentences),
                'avg_words_per_sentence': round(avg_words_per_sentence, 2),
                'avg_syllables_per_word': round(avg_syllables_per_word, 2)
            })
        
        return readability
    
    def _analyze_complexity(self, text: str) -> Dict[str, Any]:
        """Analyze text complexity."""
        words = text.split()
        
        if not words:
            return {}
        
        # Lexical diversity
        unique_words = set(word.lower() for word in words)
        lexical_diversity = len(unique_words) / len(words)
        
        # Long words (>= 7 characters)
        long_words = [word for word in words if len(word) >= 7]
        long_word_ratio = len(long_words) / len(words)
        
        # Complex punctuation
        complex_punct = len(re.findall(r'[;:(){}[\]"]', text))
        
        # Technical terms (words with numbers, capitals, etc.)
        technical_terms = len(re.findall(r'\b[A-Z]{2,}|\b\w*\d\w*\b', text))
        
        return {
            'lexical_diversity': round(lexical_diversity, 3),
            'long_word_ratio': round(long_word_ratio, 3),
            'complex_punctuation_count': complex_punct,
            'technical_terms_count': technical_terms,
            'avg_word_length': round(sum(len(word) for word in words) / len(words), 2)
        }
    
    def _extract_keywords(self, text: str, max_keywords: int = 20) -> List[str]:
        """Extract keywords from text."""
        keywords = []
        
        if self.nlp:
            try:
                # Use spaCy for advanced keyword extraction
                doc = self.nlp(text[:1000000])  # Limit text length
                
                # Extract named entities
                entities = [ent.text for ent in doc.ents if ent.label_ in ['PERSON', 'ORG', 'GPE', 'EVENT']]
                
                # Extract important nouns and adjectives
                important_words = []
                for token in doc:
                    if (token.pos_ in ['NOUN', 'PROPN', 'ADJ'] and 
                        len(token.text) > 3 and 
                        not token.is_stop and 
                        token.is_alpha and
                        not token.is_punct):
                        important_words.append(token.lemma_.lower())
                
                # Combine and deduplicate
                keywords = list(set(entities + important_words))
                
            except Exception as e:
                logger.warning(f"SpaCy keyword extraction failed: {str(e)}")
        
        # Fallback to simple keyword extraction
        if not keywords:
            keywords = self._simple_keyword_extraction(text)
        
        # Sort by frequency and limit
        word_freq = {}
        text_lower = text.lower()
        for keyword in keywords:
            word_freq[keyword] = text_lower.count(keyword.lower())
        
        keywords = sorted(keywords, key=lambda x: word_freq[x], reverse=True)
        return keywords[:max_keywords]
    
    def _simple_keyword_extraction(self, text: str) -> List[str]:
        """Simple keyword extraction using basic NLP."""
        # Remove punctuation and convert to lowercase
        clean_text = re.sub(r'[^\w\s]', ' ', text.lower())
        words = clean_text.split()
        
        # Common stop words
        stop_words = set([
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with',
            'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'do', 'does',
            'did', 'will', 'would', 'could', 'should', 'may', 'might', 'can', 'this', 'that',
            'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her',
            'us', 'them', 'my', 'your', 'his', 'her', 'its', 'our', 'their'
        ])
        
        # Filter words
        keywords = []
        for word in words:
            if (len(word) > 3 and 
                word not in stop_words and 
                word.isalpha() and
                not word.isdigit()):
                keywords.append(word)
        
        # Remove duplicates while preserving order
        seen = set()
        unique_keywords = []
        for keyword in keywords:
            if keyword not in seen:
                seen.add(keyword)
                unique_keywords.append(keyword)
        
        return unique_keywords
    
    def _extract_topics(self, text: str, max_topics: int = 5) -> List[str]:
        """Extract main topics from text."""
        topics = []
        
        # Domain-based topic detection
        topic_domains = {
            'technology': ['computer', 'software', 'digital', 'internet', 'data', 'algorithm', 'programming'],
            'science': ['research', 'study', 'analysis', 'experiment', 'theory', 'hypothesis', 'method'],
            'business': ['company', 'market', 'business', 'profit', 'customer', 'revenue', 'strategy'],
            'education': ['learning', 'teaching', 'student', 'school', 'education', 'knowledge', 'curriculum'],
            'health': ['health', 'medical', 'patient', 'treatment', 'disease', 'therapy', 'medicine'],
            'history': ['history', 'historical', 'past', 'ancient', 'century', 'period', 'era'],
            'literature': ['book', 'author', 'story', 'novel', 'literature', 'writing', 'poetry'],
            'mathematics': ['math', 'number', 'equation', 'formula', 'calculate', 'theorem', 'proof'],
            'politics': ['government', 'political', 'policy', 'law', 'election', 'democracy', 'rights'],
            'environment': ['environment', 'climate', 'nature', 'ecology', 'conservation', 'sustainability']
        }
        
        text_lower = text.lower()
        topic_scores = {}
        
        for topic, keywords in topic_domains.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            if score >= 2:  # Minimum threshold
                topic_scores[topic] = score
        
        # Sort by score and return top topics
        topics = sorted(topic_scores.keys(), key=lambda x: topic_scores[x], reverse=True)
        return topics[:max_topics]
    
    def _analyze_language_features(self, text: str) -> Dict[str, Any]:
        """Analyze language features."""
        features = {}
        
        # Question count
        questions = len(re.findall(r'\?', text))
        
        # Exclamation count
        exclamations = len(re.findall(r'!', text))
        
        # Passive voice indicators
        passive_indicators = len(re.findall(r'\b(was|were|been|being)\s+\w+ed\b', text.lower()))
        
        # Complex sentence indicators
        complex_sentences = len(re.findall(r'\b(however|moreover|furthermore|nevertheless|although|because|since|while|whereas)\b', text.lower()))
        
        # Dialogue indicators
        dialogue = len(re.findall(r'["\'].*?["\']', text))
        
        features = {
            'question_count': questions,
            'exclamation_count': exclamations,
            'passive_voice_indicators': passive_indicators,
            'complex_sentence_indicators': complex_sentences,
            'dialogue_count': dialogue,
            'has_quotes': dialogue > 0,
            'has_questions': questions > 0
        }
        
        return features
    
    def _analyze_structure(self, text: str) -> Dict[str, Any]:
        """Analyze text structure."""
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        # Detect headings (lines that are short and potentially all caps or title case)
        lines = text.split('\n')
        headings = []
        for line in lines:
            line = line.strip()
            if (line and 
                len(line) < 100 and 
                (line.isupper() or line.istitle()) and
                not line.endswith('.') and
                len(line.split()) <= 10):
                headings.append(line)
        
        # Detect lists
        list_items = len(re.findall(r'^\s*[-*•]\s+', text, re.MULTILINE))
        numbered_lists = len(re.findall(r'^\s*\d+\.\s+', text, re.MULTILINE))
        
        structure = {
            'paragraph_count': len(paragraphs),
            'sentence_count': len(sentences),
            'avg_paragraph_length': round(sum(len(p.split()) for p in paragraphs) / max(len(paragraphs), 1), 2),
            'avg_sentence_length': round(sum(len(s.split()) for s in sentences) / max(len(sentences), 1), 2),
            'heading_count': len(headings),
            'list_item_count': list_items,
            'numbered_list_count': numbered_lists,
            'has_structure': len(headings) > 0 or list_items > 0 or numbered_lists > 0
        }
        
        return structure
    
    def _calculate_difficulty_score(self, analysis: Dict[str, Any]) -> float:
        """Calculate overall difficulty score (1-10)."""
        try:
            score = 5.0  # Base score
            
            # Readability contribution (40%)
            readability = analysis.get('readability', {})
            if 'flesch_reading_ease' in readability:
                flesch = readability['flesch_reading_ease']
                if flesch >= 90:
                    readability_score = 2.0
                elif flesch >= 80:
                    readability_score = 3.0
                elif flesch >= 70:
                    readability_score = 4.0
                elif flesch >= 60:
                    readability_score = 5.0
                elif flesch >= 50:
                    readability_score = 6.0
                elif flesch >= 30:
                    readability_score = 7.0
                else:
                    readability_score = 8.0
                
                score = score * 0.6 + readability_score * 0.4
            
            # Complexity contribution (30%)
            complexity = analysis.get('complexity', {})
            if complexity:
                complexity_factors = [
                    complexity.get('lexical_diversity', 0.5),
                    complexity.get('long_word_ratio', 0.1),
                    min(complexity.get('avg_word_length', 5) / 10, 1.0)
                ]
                complexity_score = sum(complexity_factors) / len(complexity_factors) * 10
                score = score * 0.7 + complexity_score * 0.3
            
            # Language features contribution (20%)
            features = analysis.get('language_features', {})
            if features:
                feature_complexity = (
                    features.get('complex_sentence_indicators', 0) * 0.1 +
                    features.get('passive_voice_indicators', 0) * 0.05
                )
                feature_score = min(feature_complexity, 3.0) + 5.0
                score = score * 0.8 + feature_score * 0.2
            
            # Structure contribution (10%)
            structure = analysis.get('structure', {})
            if structure:
                avg_sentence_length = structure.get('avg_sentence_length', 15)
                if avg_sentence_length > 25:
                    structure_score = 7.0
                elif avg_sentence_length > 20:
                    structure_score = 6.0
                elif avg_sentence_length > 15:
                    structure_score = 5.0
                else:
                    structure_score = 4.0
                
                score = score * 0.9 + structure_score * 0.1
            
            return round(max(1.0, min(10.0, score)), 1)
            
        except Exception as e:
            logger.warning(f"Difficulty calculation failed: {str(e)}")
            return 5.0
    
    def _get_reading_level(self, flesch_score: float) -> str:
        """Convert Flesch reading ease score to reading level."""
        if flesch_score >= 90:
            return "Very Easy (5th grade)"
        elif flesch_score >= 80:
            return "Easy (6th grade)"
        elif flesch_score >= 70:
            return "Fairly Easy (7th grade)"
        elif flesch_score >= 60:
            return "Standard (8th-9th grade)"
        elif flesch_score >= 50:
            return "Fairly Difficult (10th-12th grade)"
        elif flesch_score >= 30:
            return "Difficult (College level)"
        else:
            return "Very Difficult (Graduate level)"
    
    def _count_syllables(self, word: str) -> int:
        """Count syllables in a word (simple approximation)."""
        word = word.lower()
        vowels = 'aeiouy'
        syllable_count = 0
        prev_was_vowel = False
        
        for char in word:
            is_vowel = char in vowels
            if is_vowel and not prev_was_vowel:
                syllable_count += 1
            prev_was_vowel = is_vowel
        
        # Handle silent 'e'
        if word.endswith('e') and syllable_count > 1:
            syllable_count -= 1
        
        return max(1, syllable_count)


class ContentValidator:
    """Validate content for question generation suitability."""
    
    def __init__(self):
        self.min_word_count = 50
        self.max_word_count = 50000
        self.min_sentences = 3
    
    def validate_content(self, text: str, metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """Validate content for question generation."""
        try:
            validation_result = {
                'is_valid': True,
                'warnings': [],
                'errors': [],
                'suggestions': [],
                'suitability_score': 0.0
            }
            
            if not text or not text.strip():
                validation_result['is_valid'] = False
                validation_result['errors'].append("Content is empty")
                return validation_result
            
            # Word count validation
            words = text.split()
            word_count = len(words)
            
            if word_count < self.min_word_count:
                validation_result['errors'].append(f"Content too short ({word_count} words). Minimum: {self.min_word_count}")
                validation_result['is_valid'] = False
            elif word_count > self.max_word_count:
                validation_result['warnings'].append(f"Content very long ({word_count} words). May affect processing time.")
            
            # Sentence validation
            sentences = re.split(r'[.!?]+', text)
            sentences = [s.strip() for s in sentences if s.strip()]
            
            if len(sentences) < self.min_sentences:
                validation_result['errors'].append(f"Too few sentences ({len(sentences)}). Minimum: {self.min_sentences}")
                validation_result['is_valid'] = False
            
            # Content quality checks
            self._check_content_quality(text, validation_result)
            
            # Calculate suitability score
            if validation_result['is_valid']:
                validation_result['suitability_score'] = self._calculate_suitability_score(text, validation_result)
            
            return validation_result
            
        except Exception as e:
            return {
                'is_valid': False,
                'warnings': [],
                'errors': [f"Validation failed: {str(e)}"],
                'suggestions': [],
                'suitability_score': 0.0
            }
    
    def _check_content_quality(self, text: str, validation_result: Dict[str, Any]):
        """Check content quality and add warnings/suggestions."""
        
        # Check for repeated content
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip().lower() for s in sentences if s.strip()]
        unique_sentences = set(sentences)
        
        if len(unique_sentences) < len(sentences) * 0.8:
            validation_result['warnings'].append("Content has significant repetition")
        
        # Check for gibberish or non-meaningful content
        words = text.split()
        if words:
            avg_word_length = sum(len(word) for word in words) / len(words)
            if avg_word_length < 2 or avg_word_length > 15:
                validation_result['warnings'].append("Unusual word length distribution detected")
        
        # Check for educational content indicators
        educational_indicators = [
            'learn', 'understand', 'concept', 'theory', 'principle', 'method',
            'example', 'definition', 'explanation', 'analysis', 'conclusion'
        ]
        
        text_lower = text.lower()
        indicator_count = sum(1 for indicator in educational_indicators if indicator in text_lower)
        
        if indicator_count < 2:
            validation_result['suggestions'].append("Content may benefit from more educational structure (definitions, examples, explanations)")
        
        # Check for question generation potential
        if '?' not in text and 'what' not in text_lower and 'how' not in text_lower:
            validation_result['suggestions'].append("Consider adding examples or case studies to improve question generation potential")
        
        # Check for factual content
        fact_indicators = ['research', 'study', 'data', 'evidence', 'findings', 'results']
        if any(indicator in text_lower for indicator in fact_indicators):
            validation_result['suggestions'].append("Factual content detected - good for knowledge-based questions")
        
        # Check for procedural content
        procedure_indicators = ['step', 'process', 'method', 'procedure', 'first', 'next', 'then', 'finally']
        if any(indicator in text_lower for indicator in procedure_indicators):
            validation_result['suggestions'].append("Procedural content detected - good for sequence and process questions")
    
    def _calculate_suitability_score(self, text: str, validation_result: Dict[str, Any]) -> float:
        """Calculate content suitability score for question generation (0-1)."""
        score = 0.5  # Base score
        
        # Length appropriateness (0.2)
        words = text.split()
        word_count = len(words)
        
        if 100 <= word_count <= 5000:
            length_score = 1.0
        elif 50 <= word_count < 100 or 5000 < word_count <= 10000:
            length_score = 0.8
        elif word_count < 50 or word_count > 10000:
            length_score = 0.4
        else:
            length_score = 0.6
        
        score += (length_score - 0.5) * 0.2
        
        # Structural quality (0.2)
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        if len(sentences) >= 5:
            structure_score = 1.0
        elif len(sentences) >= 3:
            structure_score = 0.8
        else:
            structure_score = 0.4
        
        score += (structure_score - 0.5) * 0.2
        
        # Educational content indicators (0.3)
        educational_keywords = [
            'definition', 'concept', 'theory', 'principle', 'method', 'process',
            'example', 'analysis', 'explanation', 'research', 'study', 'findings'
        ]
        
        text_lower = text.lower()
        educational_score = min(1.0, sum(1 for keyword in educational_keywords if keyword in text_lower) / 5)
        
        score += (educational_score - 0.5) * 0.3
        
        # Diversity and richness (0.3)
        unique_words = set(word.lower() for word in words)
        diversity = len(unique_words) / max(len(words), 1)
        
        if diversity >= 0.6:
            diversity_score = 1.0
        elif diversity >= 0.4:
            diversity_score = 0.8
        elif diversity >= 0.3:
            diversity_score = 0.6
        else:
            diversity_score = 0.4
        
        score += (diversity_score - 0.5) * 0.3
        
        # Penalty for warnings and errors
        warning_penalty = len(validation_result.get('warnings', [])) * 0.05
        error_penalty = len(validation_result.get('errors', [])) * 0.1
        
        score = max(0.0, min(1.0, score - warning_penalty - error_penalty))
        
        return round(score, 3)


# Utility functions

def get_file_info(file_path: str) -> Dict[str, Any]:
    """Get basic file information."""
    try:
        if not os.path.exists(file_path):
            return {'error': 'File not found'}
        
        stat = os.stat(file_path)
        
        return {
            'file_path': file_path,
            'file_name': os.path.basename(file_path),
            'file_size': stat.st_size,
            'file_size_mb': round(stat.st_size / (1024 * 1024), 2),
            'created_time': stat.st_ctime,
            'modified_time': stat.st_mtime,
            'file_extension': Path(file_path).suffix.lower(),
            'mime_type': mimetypes.guess_type(file_path)[0]
        }
        
    except Exception as e:
        return {'error': str(e)}


def clean_text(text: str) -> str:
    """Clean and normalize text."""
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove control characters
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
    
    # Normalize quotes
    text = re.sub(r'[""''`]', '"', text)
    
    # Normalize dashes
    text = re.sub(r'[–—]', '-', text)
    
    # Remove multiple consecutive punctuation
    text = re.sub(r'([.!?]){2,}', r'\1', text)
    
    return text.strip()


def estimate_processing_time(content_length: int, content_type: str) -> int:
    """Estimate processing time in seconds."""
    base_time = 5  # Base processing time
    
    # Time per 1000 characters
    time_per_1k = {
        'text': 0.1,
        'pdf': 0.5,
        'docx': 0.3,
        'audio': 2.0,  # Transcription is slower
        'video': 3.0,
        'url': 1.0
    }
    
    rate = time_per_1k.get(content_type, 0.5)
    estimated = base_time + (content_length / 1000) * rate
    
    return int(estimated)